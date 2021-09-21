### 參考信件(無註解)：https://groups.google.com/g/python-docx/c/0ZvpEAGH4-U
### (有註解)：https://zh.codeprj.com/blog/a456b11.html
import docx
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import RGBColor

import docx
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.dml import MSO_THEME_COLOR_INDEX

new_doc = docx.Document()
p = new_doc.add_paragraph()

def add_hyperlink(document, paragraph, url, name):
    """
    Add a hyperlink to a paragraph.

    :param document: The Document being edited.
    :param paragraph: The Paragraph the hyperlink is being added to.
    :param url: The url to be added to the link.
    :param name: The text for the link to be displayed in the paragraph
    :return: None
    """

    part = document.part
    rId = part.relate_to(url, RT.HYPERLINK, is_external=True)

    init_hyper = OxmlElement('w:hyperlink')
    init_hyper.set(qn('r:id'), rId, )
    init_hyper.set(qn('w:history'), '1')

    new_run = OxmlElement('w:r')

    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')

    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = name
    init_hyper.append(new_run)

    r = paragraph.add_run()
    r._r.append(init_hyper)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return None

def add_hyperlink_kong(document, paragraph, url, text):
    """
    Add a hyperlink to a paragraph.

    document: The Document being edited.
    paragraph: The Paragraph the hyperlink is being added to.
    url: The url to be added to the link.
    text: The text for the link to be displayed in the paragraph
    
    return: None
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = document.part
    rId = part.relate_to(url, RT.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    init_hyper = OxmlElement('w:hyperlink')
    init_hyper.set(qn('r:id'), rId, )
    init_hyper.set(qn('w:history'), '1')

    # Create a w:r element
    r = OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')

    ### 不知道
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')

    # Join all the xml elements together add add the required text to the w:r element
    rPr.append(rStyle)
    r.append(rPr)
    r.text = text
    init_hyper.append(r)

    run = paragraph.add_run()
    run._r.append(init_hyper)
    run.font.color.rgb = RGBColor(5, 99, 193)  ### 這rgb是自己截圖word預設的超連結後抓取顏色得到的rgb
    run.font.underline = True

    return run

if(__name__ == "__main__"):
    new_doc = docx.Document()
    p = new_doc.add_paragraph()

    add_hyperlink(new_doc, p, 'http://google.com', 'google')
    new_doc.save("experiment.docx")
