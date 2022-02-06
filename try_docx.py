from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm, Inches, Pt, RGBColor

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from numpy import average

from try_hyperlink import add_hyperlink_kong

import datetime
import time

class Docx_util:
    def __init__(self):
        self.docx = Document()
        self.current_time = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

    def _is_url_string(self, content):
        if( "http" in content[:5].lower()): return True
        else: return False

    def _is_img_string(self, content):
        if( ".jpg" in content[-4:].lower()): return True
        else: return False

    def _set_cell_style(self, cell, content, cell_width, bold=False, font_name="", font_size=Pt(12), font_r=0, font_g=0, font_b=0, bg_r=255, bg_g=255, bg_b=255):
        ### 設定置中：https://stackoverflow.com/questions/42963372/add-paragraph-in-docx-adds-newline
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ### 儲存格的垂直置中：https://stackoverflow.com/questions/43074163/python-docx-vertical-alignment-in-cell
        ### 設定格式：https://www.codegrepper.com/code-examples/python/python+docx+bold+heading
        run = cell.paragraphs[0].add_run()
        if(self._is_img_string(content) is False):
            if(self._is_url_string(content) is True):
                run = add_hyperlink_kong(self.docx, cell.paragraphs[0], url=content, text=content)
            else: run.text           = content

            run.bold           = bold
            run.font.name      = font_name
            run.font.size      = font_size
            run.font.color.rgb = RGBColor(font_r, font_g, font_b)

        else:
            
            run.add_picture(content, height=Pt(12) * 4.5)   ### 沒辦法就要慢慢嘗試囉～～
            # run.add_picture(content, width=cell.width)

        ### 設定單格cell顏色，從這邊複製的：https://groups.google.com/g/python-docx/c/-c3OrRHA3qo
        bg_color = "%s%s%s" % (hex(bg_r)[2:], hex(bg_g)[2:], hex(bg_b)[2:])
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        ### 想print看看是什麼，結果也看不懂，但還是留一下好了～
        # print("nsdecls('w'):", repr(nsdecls('w')))
        # print("shading_elm:", shading_elm)
        # print("cell._tc:", cell._tc)
        # print("cell._tc.get_or_add_tcPr():", cell._tc.get_or_add_tcPr())

        cell.width = cell_width

    def first_r_style(self, cell, content, cell_width): 
        self._set_cell_style    (cell, content, cell_width, bold=True, font_size=Pt(11), font_r=0, font_g=0, font_b=0, bg_r=180, bg_g=180, bg_b=180)
    def title_style  (self, cell, content, cell_width):
        self._set_cell_style    (cell, content, cell_width, bold=True, font_size=Pt(12), font_r=0, font_g=0, font_b=0, bg_r=255, bg_g=217, bg_b=102)
    def prod_style   (self, cell, content, cell_width, font_size=Pt(11), font_r=0, font_g=0, font_b=0):
        self._set_cell_style    (cell, content, cell_width, bold=False, font_name="新細明體", font_size=font_size, font_r=font_r, font_g=font_g, font_b=font_b)

    def set_boarder(self):
        section = self.docx.sections[0]
        ### 調整文件左右上下邊界至 1 cm，參考超強網頁：https://hackmd.io/@amostsai/SJx7_J-xN?type=view
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        return section

    def try_mogon(self, dst_dir="try_word", search_obj=None):
        import os
        os.makedirs(f'{dst_dir}', exist_ok=True)
        ################################################################################################################################
        section = self.set_boarder()
        average_cell_width = (section.page_width - Cm(1) * 2) / 4

        table = self.docx.add_table(rows=1, cols=4)  ### p115, https://readthedocs.org/projects/python-docx/downloads/pdf/stable/
        table.style = 'Table Grid'  ### 表格樣式，相當於word裡面進 設計→表格樣式 挑選的東西，可參考：https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
        table.autofit = False

        ### 第一行
        first_r_cols  = table.rows[0].cells
        self.first_r_style(cell=first_r_cols[0], content="Product_name",  cell_width=average_cell_width)
        self.first_r_style(cell=first_r_cols[1], content="Product_price", cell_width=average_cell_width)
        self.first_r_style(cell=first_r_cols[2], content="Product_link",  cell_width=average_cell_width)
        self.first_r_style(cell=first_r_cols[3], content="Product_image", cell_width=average_cell_width)

        ### products
        for go_prod, prod in enumerate(search_obj.containor.prods):
            ### series
            if( (go_prod == 0 and prod.series is not None) or 
                (go_prod != 0 and prod.series != search_obj.containor.prods[go_prod - 1].series)):
                series_r_cols = table.add_row().cells
                series_r_cols[0].merge(series_r_cols[-1])  ### 合併儲存格 p126, https://readthedocs.org/projects/python-docx/downloads/pdf/stable/
                self.title_style(series_r_cols[0], prod.series, average_cell_width)

            ### product
            prod_r_cols = table.add_row().cells
            self.prod_style(cell=prod_r_cols[0], content=prod.prod_title,                         cell_width=average_cell_width)
            self.prod_style(cell=prod_r_cols[1], content=prod.price,                              cell_width=average_cell_width, font_size=Pt(14))
            self.prod_style(cell=prod_r_cols[2], content=prod.prod_url,                           cell_width=average_cell_width, font_size=Pt(8), font_r=5, font_g=99, font_b=193)
            self.prod_style(cell=prod_r_cols[3], content=f"{dst_dir}/imgs/"+"%04i.jpg" % (go_prod + 1), cell_width=average_cell_width)
        self.docx.save(f'{dst_dir}/{self.current_time}.docx')

    
    def try_docx(self, dst_dir="try_word"):
        import os
        
        os.makedirs(f'{dst_dir}', exist_ok=True)
        ################################################################################################################################
        # docx = Document()
        section = self.set_boarder()
        average_cell_width = (section.page_width - Cm(1) * 2) / 4
        # print("average_cell_width",  average_cell_width)
        # print("section.page_width",  section.page_width)
        # print("section.page_height", section.page_height)


        table = self.docx.add_table(rows=1, cols=4)  ### p115, https://readthedocs.org/projects/python-docx/downloads/pdf/stable/
        table.style = 'Table Grid'  ### 表格樣式，相當於word裡面進 設計→表格樣式 挑選的東西，可參考：https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
        table.autofit = False

        ### 第一行
        first_r_cols  = table.rows[0].cells
        self.first_r_style(cell=first_r_cols[0], content="Product_name",  cell_width=average_cell_width)
        self.first_r_style(cell=first_r_cols[1], content="Product_price", cell_width=average_cell_width)
        self.first_r_style(cell=first_r_cols[2], content="Product_link",  cell_width=average_cell_width)
        self.first_r_style(cell=first_r_cols[3], content="Product_image", cell_width=average_cell_width)


        ################################################################################################################################
        for _ in range(5):
            ### title
            series_r_cols = table.add_row().cells
            series_r_cols[0].merge(series_r_cols[-1])  ### 合併儲存格 p126, https://readthedocs.org/projects/python-docx/downloads/pdf/stable/
            self.title_style(series_r_cols[0], "M-1000 Series", average_cell_width)
            for _ in range(2):
                ### product
                prod_r_cols = table.add_row().cells
                self.prod_style(cell=prod_r_cols[0], content="PlayWood/プレイウッド　吉岡孝悦モデル マリンバ用キーボードマレット　M-1001R", cell_width=average_cell_width)
                self.prod_style(cell=prod_r_cols[1], content="1,321 NTD", cell_width=average_cell_width, font_size=Pt(14))
                self.prod_style(cell=prod_r_cols[2], content="https://www.moganshopping.com/zh_tw/public/jyahooshopping/auctionItem.php?Code=rizing_pw-m-1001r&categoryid=1", cell_width=average_cell_width, font_size=Pt(8), font_r=5, font_g=99, font_b=193)
                self.prod_style(cell=prod_r_cols[3], content="try_0001.jpg", cell_width=average_cell_width)


        ################################################################################################################################
        self.docx.save(f'{dst_dir}/{self.current_time}.docx')
        print("save to:", f'{dst_dir}/try_word.docx')

if(__name__ == "__main__"):
    
    dst_dir = f"try_word"
    try_word = Docx_util()
    try_word.try_docx(dst_dir=dst_dir)
    
    ###########################################################################################################
    open_docx_keep_time = 10
    ### 自動打開看結果，等三秒後自動關閉
    from win32com import client
    word = client.gencache.EnsureDispatch('word.application')
    # word.Documents.Open(f"C:/Users/TKU/Desktop/m/{dst_dir}/{current_time}.docx")  ### 一定要絕對位置
    word.Documents.Open(f"C:/Users/HP820G1/Desktop/kong_mogon/{dst_dir}/{try_word.current_time}.docx")  ### 一定要絕對位置
    word.Visible = 1
    word.DisplayAlerts = 0
    time.sleep(open_docx_keep_time)
    word.Quit()
